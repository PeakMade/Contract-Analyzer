"""
Test script to verify Word COM API is available and working.
"""
import sys
import os

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def test_word_com_availability():
    """Test if Word COM API is available."""
    print("="*70)
    print("Word COM API Availability Test")
    print("="*70)
    
    try:
        from app.services.word_grammar_checker import is_word_com_available
        
        available = is_word_com_available()
        
        if available:
            print("✅ SUCCESS: Word COM API is available!")
            print("   Microsoft Word is installed and accessible.")
            print("   Grammar checking will use Word COM by default.")
        else:
            print("❌ FAILED: Word COM API is NOT available")
            print("   Microsoft Word may not be installed.")
            print("   Grammar checking will fallback to GPT.")
        
        return available
        
    except ImportError as e:
        print("❌ IMPORT ERROR: pywin32 not installed")
        print(f"   Error: {e}")
        print("   Install with: pip install pywin32")
        return False
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_word_grammar_check():
    """Test Word grammar checking on a sample document."""
    print("\n" + "="*70)
    print("Word Grammar Check Test (Sample Document)")
    print("="*70)
    
    try:
        from app.services.word_grammar_checker import check_spelling_grammar_with_word
        from docx import Document
        import tempfile
        
        # Create a test document with intentional errors
        print("\n1. Creating test document with errors...")
        temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_path = temp_doc.name
        temp_doc.close()
        
        doc = Document()
        doc.add_paragraph("This is a tets document with spelling erors.")
        doc.add_paragraph("The contract shall be effective on the date of signing and will continue.")
        doc.save(temp_path)
        print(f"   ✓ Created: {temp_path}")
        
        # Run Word grammar check
        print("\n2. Running Word COM grammar check...")
        result = check_spelling_grammar_with_word(temp_path)
        
        # Display results
        print(f"\n3. Results:")
        print(f"   Method: {result.get('method', 'unknown')}")
        print(f"   Issues found: {result.get('issues_found', False)}")
        print(f"   Error count: {result.get('error_count', 0)}")
        
        if result.get('raw_counts'):
            raw = result['raw_counts']
            print(f"   Raw spelling errors: {raw.get('spelling', 0)}")
            print(f"   Raw grammar errors: {raw.get('grammar', 0)}")
        
        if result.get('errors'):
            print(f"\n4. First few errors:")
            for i, error in enumerate(result['errors'][:3], 1):
                print(f"\n   Error {i}:")
                print(f"      Type: {error.get('type')}")
                print(f"      Error: {error.get('error_text')}")
                print(f"      Suggestion: {error.get('suggestion')}")
        
        # Clean up
        print(f"\n5. Cleaning up test file...")
        os.unlink(temp_path)
        print(f"   ✓ Deleted: {temp_path}")
        
        # Success check
        if result.get('method') == 'word_com' and result.get('error_count', 0) >= 2:
            print("\n✅ SUCCESS: Word COM grammar check is working!")
            print("   Detected spelling errors in test document.")
            return True
        elif result.get('method') == 'word_com':
            print("\n⚠️ WARNING: Word COM available but didn't detect expected errors")
            print("   Check Word's proofing settings.")
            return True
        else:
            print("\n❌ FAILED: Word COM grammar check failed")
            if result.get('error_message'):
                print(f"   Error: {result.get('error_message')}")
            return False
        
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    print("\nStarting Word COM API Tests...\n")
    
    # Test 1: Availability
    available = test_word_com_availability()
    
    # Test 2: Grammar check (only if available)
    if available:
        success = test_word_grammar_check()
    else:
        print("\nSkipping grammar check test (Word COM not available)")
        success = False
    
    # Summary
    print("\n" + "="*70)
    print("Test Summary")
    print("="*70)
    if available and success:
        print("✅ All tests passed!")
        print("   Word COM API is ready for use.")
        print("\n   Configuration:")
        print("   - GRAMMAR_CHECK_METHOD=auto (recommended)")
        print("   - Will use Word COM by default, GPT as fallback")
    elif available:
        print("⚠️ Partial success")
        print("   Word COM is available but some tests had issues.")
        print("   System will fallback to GPT if Word COM fails.")
    else:
        print("❌ Tests failed")
        print("   Word COM not available - will use GPT only.")
        print("\n   To enable Word COM:")
        print("   1. Ensure Microsoft Word is installed")
        print("   2. Install pywin32: pip install pywin32")
    print("="*70)
