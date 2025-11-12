"""
Unit tests for Apply Selected Suggestions feature.
Tests doc_editor, sp_upload, and route integration.
"""

import unittest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path
import tempfile
import json
from docx import Document
from io import BytesIO

from app.services import doc_editor, sp_upload


class TestStyleDetector(unittest.TestCase):
    """Test style detection in documents."""
    
    def setUp(self):
        """Create test document."""
        self.doc = Document()
        self.known_standards = ['Payment Terms', 'Termination', 'Confidentiality']
    
    def test_detect_heading_style_found(self):
        """Should detect heading style when standard is found."""
        # Add paragraph with standard name using Heading 3
        para = self.doc.add_paragraph('Payment Terms', style='Heading 3')
        
        detector = doc_editor.StyleDetector(self.doc, self.known_standards)
        style = detector.detect_heading_style()
        
        self.assertEqual(style, 'Heading 3')
    
    def test_detect_heading_style_fallback(self):
        """Should use fallback when no standard found."""
        # Add paragraphs without standard names
        self.doc.add_paragraph('Some random text')
        
        detector = doc_editor.StyleDetector(self.doc, self.known_standards)
        style = detector.detect_heading_style()
        
        self.assertEqual(style, 'Heading 2')
    
    def test_detect_body_style_found(self):
        """Should detect body style following standard heading."""
        # Add standard heading
        self.doc.add_paragraph('Termination', style='Heading 2')
        # Add body paragraph
        self.doc.add_paragraph('This is the body text.', style='Body Text')
        
        detector = doc_editor.StyleDetector(self.doc, self.known_standards)
        style = detector.detect_body_style()
        
        self.assertEqual(style, 'Body Text')
    
    def test_detect_body_style_fallback(self):
        """Should use fallback when no body found."""
        detector = doc_editor.StyleDetector(self.doc, self.known_standards)
        style = detector.detect_body_style()
        
        self.assertEqual(style, 'Normal')
    
    def test_case_insensitive_detection(self):
        """Should detect styles case-insensitively."""
        self.doc.add_paragraph('payment terms', style='Heading 4')
        
        detector = doc_editor.StyleDetector(self.doc, self.known_standards)
        style = detector.detect_heading_style()
        
        self.assertEqual(style, 'Heading 4')


class TestAppendSuggestedStandards(unittest.TestCase):
    """Test document editing functionality."""
    
    def setUp(self):
        """Create temporary test document."""
        self.temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        self.temp_path = Path(self.temp_doc.name)
        self.temp_doc.close()
        
        # Create a simple document
        doc = Document()
        doc.add_heading('Existing Contract', level=1)
        doc.add_paragraph('This is existing content.')
        doc.save(str(self.temp_path))
        
        self.items = [
            {'standard': 'Payment Terms', 'suggestion': 'Net 30 payment required.'},
            {'standard': 'Termination', 'suggestion': 'Either party may terminate with 30 days notice.'}
        ]
    
    def tearDown(self):
        """Clean up temp files."""
        if self.temp_path.exists():
            self.temp_path.unlink()
    
    def test_append_creates_new_file(self):
        """Should create edited document."""
        result_path = doc_editor.append_suggested_standards(
            self.temp_path,
            self.items
        )
        
        self.assertTrue(result_path.exists())
        self.assertNotEqual(result_path, self.temp_path)
        
        # Cleanup
        result_path.unlink()
    
    def test_append_adds_appendix_heading(self):
        """Should add 'Appendix — Suggested Standards' heading."""
        result_path = doc_editor.append_suggested_standards(
            self.temp_path,
            self.items
        )
        
        doc = Document(str(result_path))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        self.assertIn('Appendix — Suggested Standards', headings)
        
        # Cleanup
        result_path.unlink()
    
    def test_append_adds_all_standards(self):
        """Should add all provided standards."""
        result_path = doc_editor.append_suggested_standards(
            self.temp_path,
            self.items
        )
        
        doc = Document(str(result_path))
        text = ' '.join([p.text for p in doc.paragraphs])
        
        for item in self.items:
            self.assertIn(item['standard'], text)
            self.assertIn(item['suggestion'], text)
        
        # Cleanup
        result_path.unlink()
    
    def test_file_not_found_raises(self):
        """Should raise FileNotFoundError for missing file."""
        fake_path = Path('/nonexistent/file.docx')
        
        with self.assertRaises(FileNotFoundError):
            doc_editor.append_suggested_standards(fake_path, self.items)
    
    def test_empty_items_raises(self):
        """Should raise ValueError for empty items."""
        with self.assertRaises(ValueError):
            doc_editor.append_suggested_standards(self.temp_path, [])
    
    def test_malformed_items_raises(self):
        """Should raise ValueError for items missing keys."""
        bad_items = [{'standard': 'Test'}]  # Missing 'suggestion'
        
        with self.assertRaises(ValueError):
            doc_editor.append_suggested_standards(self.temp_path, bad_items)


class TestSpUpload(unittest.TestCase):
    """Test SharePoint upload functionality."""
    
    def setUp(self):
        """Set up Flask test client for request context."""
        from main import app
        self.app = app
        self.app.config['TESTING'] = True
    
    @patch('app.services.sp_upload.requests.put')
    def test_upload_success(self, mock_put):
        """Should upload file successfully."""
        with self.app.test_request_context():
            from flask import session
            session['access_token'] = 'fake_token'
            
            mock_response = Mock()
            mock_response.status_code = 201
            mock_response.json.return_value = {'id': '123', 'name': 'test.docx'}
            mock_put.return_value = mock_response
            
            result = sp_upload.upload_file(
                drive_id='drive123',
                folder_path='',
                filename='test.docx',
                content=b'fake content'
            )
            
            self.assertEqual(result['id'], '123')
            self.assertEqual(result['name'], 'test.docx')
    
    def test_upload_no_token_raises(self):
        """Should raise PermissionError when no token."""
        with self.app.test_request_context():
            with self.assertRaises(PermissionError) as ctx:
                sp_upload.upload_file('drive123', '', 'test.docx', b'content')
            
            self.assertIn('SESSION_EXPIRED', str(ctx.exception))
    
    @patch('app.services.sp_upload.requests.put')
    def test_upload_401_raises_permission_error(self, mock_put):
        """Should raise PermissionError on 401 response."""
        with self.app.test_request_context():
            from flask import session
            session['access_token'] = 'fake_token'
            
            mock_response = Mock()
            mock_response.status_code = 401
            mock_put.return_value = mock_response
            
            with self.assertRaises(PermissionError):
                sp_upload.upload_file('drive123', '', 'test.docx', b'content')
    
    @patch('app.services.sp_upload.requests.put')
    def test_upload_other_error_raises_upload_error(self, mock_put):
        """Should raise UploadError on non-401 errors."""
        with self.app.test_request_context():
            from flask import session
            session['access_token'] = 'fake_token'
            
            mock_response = Mock()
            mock_response.status_code = 500
            mock_response.text = 'Server error'
            mock_response.json.side_effect = ValueError()
            mock_put.return_value = mock_response
            
            with self.assertRaises(sp_upload.UploadError):
                sp_upload.upload_file('drive123', '', 'test.docx', b'content')
    
    def test_generate_edited_filename(self):
        """Should add _edited suffix correctly."""
        self.assertEqual(
            sp_upload.generate_edited_filename('Contract.docx'),
            'Contract_edited.docx'
        )
        self.assertEqual(
            sp_upload.generate_edited_filename('My_File.DOCX'),
            'My_File_edited.DOCX'
        )
        self.assertEqual(
            sp_upload.generate_edited_filename('noext'),
            'noext_edited'
        )


class TestApplySuggestionsRoute(unittest.TestCase):
    """Test the apply_suggestions route."""
    
    def setUp(self):
        """Set up Flask test client."""
        from main import app
        self.app = app
        self.client = app.test_client()
        self.app.config['TESTING'] = True
    
    @unittest.skip("Integration test - requires real SharePoint connection")
    def test_apply_suggestions_integration(self):
        """
        Full integration test for apply_suggestions workflow.
        Skipped in unit tests - run manually with real SharePoint data.
        
        To run: Remove @skip decorator and provide valid contract_id
        """
        with self.client.session_transaction() as sess:
            sess['access_token'] = 'real_token_here'
            sess['user_email'] = 'test@example.com'
        
        response = self.client.post(
            '/contracts/REAL_CONTRACT_ID/apply_suggestions',
            json={
                'items': [
                    {'standard': 'Payment', 'suggestion': 'Net 30'}
                ]
            }
        )
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.data)
        self.assertTrue(data['success'])
        self.assertIn('download_url', data)
    
    def test_apply_suggestions_missing_items(self):
        """Should return 400 for missing items."""
        with self.client.session_transaction() as sess:
            sess['access_token'] = 'fake_token'
            sess['user_email'] = 'test@example.com'
        
        response = self.client.post(
            '/contracts/123/apply_suggestions',
            json={}
        )
        
        self.assertEqual(response.status_code, 400)
        data = json.loads(response.data)
        self.assertIn('error', data)
    
    @patch('app.services.sharepoint_service.requests.post')
    def test_apply_suggestions_contract_not_found(self, mock_post):
        """Should return 404 for missing contract."""
        with self.client.session_transaction() as sess:
            sess['access_token'] = 'fake_token'
            sess['user_email'] = 'test@example.com'
        
        # Mock SharePoint API response with no results
        mock_sp_response = Mock()
        mock_sp_response.status_code = 200
        mock_sp_response.json.return_value = {'value': []}
        mock_post.return_value = mock_sp_response
        
        response = self.client.post(
            '/contracts/999/apply_suggestions',
            json={'items': [{'standard': 'Test', 'suggestion': 'Test'}]}
        )
        
        self.assertEqual(response.status_code, 404)


if __name__ == '__main__':
    unittest.main()
