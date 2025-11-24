"""
Text extraction service for contract documents.
Supports DOCX and PDF formats.
"""
import re
import logging
from pathlib import Path
from typing import Optional, Dict, List
import docx
from docx.oxml.ns import qn
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfparser import PDFSyntaxError

logger = logging.getLogger(__name__)

# Maximum characters to extract to avoid runaway prompts
MAX_TEXT_LENGTH = 2_000_000


def _normalize_whitespace(text: str) -> str:
    """
    Normalize whitespace and remove control characters.
    
    Args:
        text: Raw extracted text.
    
    Returns:
        Cleaned text with normalized whitespace.
    """
    # Remove control characters except newline, tab, carriage return
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Normalize line breaks
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Replace multiple spaces with single space (but preserve newlines)
    text = re.sub(r' +', ' ', text)
    
    # Replace multiple newlines with maximum of 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def _get_numbering_definitions(doc) -> Dict:
    """
    Extract numbering definitions from the document's numbering.xml.
    
    Args:
        doc: The docx.Document object.
    
    Returns:
        Dictionary mapping (numId, level) to formatting information.
    """
    numbering_dict = {}
    
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            logger.debug("Document has no numbering part")
            return numbering_dict
        
        # Parse numbering definitions
        numbering_element = numbering_part.element
        
        # Get all num elements (instances of numbering)
        num_elements = numbering_element.findall(qn('w:num'), numbering_element.nsmap)
        
        for num_element in num_elements:
            # numId is an ATTRIBUTE of the <w:num> element, not a child element
            num_id_attr = num_element.get(qn('w:numId'))
            if num_id_attr is None:
                continue
            
            num_id = int(num_id_attr)
            
            # Get the abstract numbering ID (this IS a child element)
            abstract_num_id_elem = num_element.find(qn('w:abstractNumId'), num_element.nsmap)
            if abstract_num_id_elem is None:
                continue
            
            abstract_num_id = int(abstract_num_id_elem.get(qn('w:val')))
            
            # Find the abstract numbering definition
            abstract_num_elements = numbering_element.findall(qn('w:abstractNum'), numbering_element.nsmap)
            
            for abstract_num in abstract_num_elements:
                abstract_num_id_attr = abstract_num.get(qn('w:abstractNumId'))
                if abstract_num_id_attr is None or int(abstract_num_id_attr) != abstract_num_id:
                    continue
                
                # Get level definitions
                lvl_elements = abstract_num.findall(qn('w:lvl'), abstract_num.nsmap)
                
                for lvl in lvl_elements:
                    lvl_attr = lvl.get(qn('w:ilvl'))
                    if lvl_attr is None:
                        continue
                    
                    level = int(lvl_attr)
                    
                    # Get the level text format (e.g., "%1.", "%1.%2")
                    lvl_text_elem = lvl.find(qn('w:lvlText'), lvl.nsmap)
                    lvl_text = lvl_text_elem.get(qn('w:val')) if lvl_text_elem is not None else None
                    
                    # Get the number format (decimal, lowerLetter, upperRoman, etc.)
                    num_fmt_elem = lvl.find(qn('w:numFmt'), lvl.nsmap)
                    num_fmt = num_fmt_elem.get(qn('w:val')) if num_fmt_elem is not None else 'decimal'
                    
                    # Store the formatting info
                    key = (num_id, level)
                    numbering_dict[key] = {
                        'lvlText': lvl_text,
                        'numFmt': num_fmt
                    }
        
        logger.debug(f"Extracted {len(numbering_dict)} numbering definitions")
        
    except Exception as e:
        logger.warning(f"Could not extract numbering definitions: {e}")
    
    return numbering_dict


def _format_number(value: int, num_fmt: str) -> str:
    """
    Format a number according to Word numbering format.
    
    Args:
        value: The number value (1, 2, 3, ...)
        num_fmt: Format type (decimal, lowerLetter, upperRoman, etc.)
    
    Returns:
        Formatted number string.
    """
    if num_fmt == 'decimal':
        return str(value)
    elif num_fmt == 'lowerLetter':
        # a, b, c, ... z, aa, ab, ...
        result = ""
        value -= 1  # 0-indexed for letters
        while value >= 0:
            result = chr(ord('a') + (value % 26)) + result
            value = value // 26 - 1
        return result
    elif num_fmt == 'upperLetter':
        result = ""
        value -= 1
        while value >= 0:
            result = chr(ord('A') + (value % 26)) + result
            value = value // 26 - 1
        return result
    elif num_fmt == 'lowerRoman':
        return _to_roman(value).lower()
    elif num_fmt == 'upperRoman':
        return _to_roman(value)
    else:
        # Default to decimal for unknown formats
        return str(value)


def _to_roman(num: int) -> str:
    """Convert an integer to a Roman numeral."""
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num


def _get_paragraph_number(para, doc, numbering_dict: Dict, counter_state: Dict) -> Optional[str]:
    """
    Extract and format the numbering prefix for a paragraph.
    
    Args:
        para: The paragraph object.
        doc: The document object.
        numbering_dict: Dictionary of numbering definitions.
        counter_state: Dictionary tracking current counter values per (numId, level).
    
    Returns:
        Formatted numbering string (e.g., "5.", "7.2", "a)") or None if not numbered.
    """
    try:
        # Access paragraph properties
        pPr = para._p.pPr
        if pPr is None:
            return None
        
        numPr = pPr.numPr
        if numPr is None:
            return None
        
        # Get numbering ID and level
        ilvl_elem = numPr.ilvl
        numId_elem = numPr.numId
        
        if ilvl_elem is None or numId_elem is None:
            return None
        
        level = ilvl_elem.val
        num_id = numId_elem.val
        
        if level is None or num_id is None:
            return None
        
        # Increment the counter for this numbering level
        counter_key = (num_id, level)
        
        # Initialize counter if not exists
        if counter_key not in counter_state:
            counter_state[counter_key] = 1
        else:
            counter_state[counter_key] += 1
        
        current_value = counter_state[counter_key]
        
        # Get the formatting for this numbering level
        formatting = numbering_dict.get(counter_key)
        
        if formatting is None:
            # Fallback: return a simple numbered format
            logger.debug(f"No formatting found for numId={num_id}, level={level}, using fallback")
            return f"{current_value}."
        
        lvl_text = formatting.get('lvlText', '%1.')
        num_fmt = formatting.get('numFmt', 'decimal')
        
        # Format the number according to the format type
        formatted_num = _format_number(current_value, num_fmt)
        
        # Replace the placeholder in lvlText (e.g., "%1" with the formatted number)
        # For multi-level numbering like "%1.%2", we'd need to track parent levels
        # For now, handle single-level (%1)
        if lvl_text:
            result = lvl_text.replace(f'%{level + 1}', formatted_num)
            return result
        else:
            return f"{formatted_num}."
    
    except Exception as e:
        logger.debug(f"Could not extract paragraph numbering: {e}")
        return None


def _extract_docx_text(path: Path) -> str:
    """
    Extract text from a DOCX file, preserving paragraph numbering.
    
    Args:
        path: Path to the DOCX file.
    
    Returns:
        Extracted text content with numbering preserved.
    
    Raises:
        RuntimeError: On extraction failure.
    """
    try:
        doc = docx.Document(path)
        
        # === ENHANCED DEBUGGING ===
        print(f"\n{'='*70}")
        print(f"TEXT EXTRACTION DEBUG: {path.name}")
        print(f"{'='*70}")
        
        # Extract numbering definitions
        print(f"[1/5] Extracting numbering definitions...")
        numbering_dict = _get_numbering_definitions(doc)
        print(f"      Found {len(numbering_dict)} numbering definitions")
        
        if numbering_dict:
            print(f"      Numbering formats:")
            for (num_id, level), fmt in list(numbering_dict.items())[:5]:
                print(f"        - numId={num_id}, level={level}: {fmt.get('lvlText', 'N/A')} ({fmt.get('numFmt', 'N/A')})")
        else:
            print(f"      ⚠ No numbering definitions found in document")
            print(f"        This means the document either:")
            print(f"        • Uses manually typed numbers (not automatic numbering)")
            print(f"        • Has no numbered sections")
        
        # Track counter state for numbering (increments as we process paragraphs)
        counter_state = {}
        
        # Extract text from all paragraphs with numbering
        print(f"\n[2/5] Processing paragraphs...")
        paragraphs = []
        numbered_para_count = 0
        total_para_count = 0
        
        # Track heading counters for outline numbering
        heading_counters = [0] * 10  # Support up to 10 heading levels
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            total_para_count += 1
            
            # Try to get paragraph numbering (list-based)
            num_text = _get_paragraph_number(para, doc, numbering_dict, counter_state)
            
            # If no list numbering, check for outline numbering (heading-based)
            if not num_text and para.style:
                style_name = para.style.name.lower()
                
                if 'heading' in style_name:
                    # Extract heading level from style name
                    import re
                    match = re.search(r'heading\s*(\d+)', style_name)
                    if match:
                        level = int(match.group(1))
                        
                        # Update counters (increment this level, reset lower levels)
                        heading_counters[level - 1] += 1
                        for j in range(level, 10):
                            heading_counters[j] = 0
                        
                        # Build number string based on level
                        if level == 1:
                            num_text = f"{heading_counters[0]}."
                        elif level == 2:
                            num_text = f"{heading_counters[0]}.{heading_counters[1]}"
                        elif level == 3:
                            num_text = f"{heading_counters[0]}.{heading_counters[1]}.{heading_counters[2]}"
                        else:
                            # Generic multi-level numbering
                            parts = [str(heading_counters[j]) for j in range(level) if heading_counters[j] > 0]
                            if parts:
                                num_text = '.'.join(parts) + '.'
            
            if num_text:
                # Prepend the numbering to the paragraph text
                full_text = f"{num_text} {para.text}"
                numbered_para_count += 1
                
                # Log first 10 numbered paragraphs for debugging
                if numbered_para_count <= 10:
                    preview = para.text[:60] + "..." if len(para.text) > 60 else para.text
                    print(f"      ✓ Para {total_para_count}: {num_text} {preview}")
                
                logger.debug(f"Numbered paragraph: {num_text} {para.text[:50]}...")
            else:
                full_text = para.text
            
            paragraphs.append(full_text)
        
        # Extract text from tables
        print(f"\n[3/5] Processing tables...")
        table_cell_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
                        table_cell_count += 1
        print(f"      Found {table_cell_count} table cells with content")
        
        text = '\n\n'.join(paragraphs)
        
        if not text.strip():
            raise RuntimeError("Document appears to be empty")
        
        # === ENHANCED DEBUGGING OUTPUT ===
        print(f"\n[4/5] Extraction Summary:")
        print(f"      Total paragraphs: {total_para_count}")
        print(f"      Numbered paragraphs: {numbered_para_count}")
        print(f"      Table cells: {table_cell_count}")
        print(f"      Total characters: {len(text)}")
        print(f"      Numbering success rate: {numbered_para_count}/{total_para_count} ({100*numbered_para_count/total_para_count:.1f}%)" if total_para_count > 0 else "      No paragraphs processed")
        
        print(f"\n[5/5] Checking for common numbering patterns in extracted text:")
        patterns = {
            "1.": text.count("1."),
            "2.": text.count("2."),
            "3.": text.count("3."),
            "I.": text.count("I."),
            "II.": text.count("II."),
            "Section": text.count("Section"),
            "Article": text.count("Article"),
        }
        
        for pattern, count in patterns.items():
            status = "✓" if count > 0 else "✗"
            print(f"      {status} '{pattern}' appears {count} times")
        
        # Save first 3000 chars to a debug file for inspection
        debug_file = Path("DEBUG_EXTRACTED_TEXT.txt")
        debug_content = f"FILE: {path.name}\n{'='*70}\n\n"
        debug_content += f"EXTRACTION STATS:\n"
        debug_content += f"  Total paragraphs: {total_para_count}\n"
        debug_content += f"  Numbered paragraphs: {numbered_para_count}\n"
        debug_content += f"  Total characters: {len(text)}\n\n"
        debug_content += f"FIRST 3000 CHARACTERS:\n{'='*70}\n"
        debug_content += text[:3000]
        debug_content += f"\n\n{'='*70}\n"
        debug_content += f"FULL TEXT ({len(text)} chars):\n{'='*70}\n\n"
        debug_content += text
        
        debug_file.write_text(debug_content, encoding='utf-8')
        print(f"\n      💾 Saved full extraction to: {debug_file.absolute()}")
        print(f"         Open this file to see exactly what the AI receives")
        
        print(f"{'='*70}\n")
        
        logger.info(f"Extracted {len(text)} characters from DOCX file ({numbered_para_count} numbered paragraphs)")
        return text
        
    except RuntimeError:
        # Re-raise RuntimeError as-is (these are intentional errors)
        raise
    except Exception as e:
        # For any other exception, log and raise a user-friendly error
        logger.error(f"Failed to extract text from DOCX: {type(e).__name__} - {str(e)}")
        raise RuntimeError("Failed to extract text from document. The file may be corrupted or in an unsupported format.")


def _extract_pdf_text(path: Path) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        path: Path to the PDF file.
    
    Returns:
        Extracted text content.
    
    Raises:
        RuntimeError: On extraction failure.
    """
    try:
        text = pdf_extract_text(str(path))
        
        if not text or not text.strip():
            raise RuntimeError("PDF appears to be empty or contains only images")
        
        logger.info(f"Extracted {len(text)} characters from PDF file")
        return text
        
    except PDFSyntaxError:
        logger.error("PDF file has syntax errors")
        raise RuntimeError("Failed to extract text from PDF. The file may be corrupted.")
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {type(e).__name__}")
        raise RuntimeError("Failed to extract text from PDF. The file may be encrypted, corrupted, or in an unsupported format.")


def extract_text(path: Path) -> str:
    """
    Extract text from a contract document (DOCX or PDF).
    
    Args:
        path: Path to the document file.
    
    Returns:
        Normalized text content with whitespace cleaned and control characters removed.
    
    Raises:
        RuntimeError: If extraction fails or file format is unsupported.
    """
    if not path.exists():
        raise RuntimeError("Contract file not found")
    
    # Determine file type by extension
    suffix = path.suffix.lower()
    
    try:
        if suffix == '.docx':
            raw_text = _extract_docx_text(path)
        elif suffix == '.pdf':
            raw_text = _extract_pdf_text(path)
        else:
            logger.error(f"Unsupported file format: {suffix}")
            raise RuntimeError(f"Unsupported file format: {suffix}. Only DOCX and PDF files are supported.")
        
        # Normalize whitespace
        normalized_text = _normalize_whitespace(raw_text)
        
        # Clamp to maximum length
        if len(normalized_text) > MAX_TEXT_LENGTH:
            logger.warning(f"Text length {len(normalized_text)} exceeds maximum {MAX_TEXT_LENGTH}, truncating")
            normalized_text = normalized_text[:MAX_TEXT_LENGTH]
        
        logger.info(f"Text extraction complete: {len(normalized_text)} characters after normalization")
        
        return normalized_text
        
    except RuntimeError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error during text extraction: {type(e).__name__}")
        raise RuntimeError("An unexpected error occurred while extracting text from the document")
