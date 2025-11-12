"""
Document editor service for appending suggested standards to DOCX files.
Detects existing styles and applies them to new content.
"""

from pathlib import Path
from typing import List, Dict, Optional
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

print("\n=== DEBUG: doc_editor.py module loaded ===")
print(f"python-docx imported successfully")


class StyleDetector:
    """Detects heading and body styles used for standards in a document."""
    
    def __init__(self, doc: Document, known_standards: List[str]):
        """
        Args:
            doc: python-docx Document object
            known_standards: List of standard names to search for (e.g., ["Payment Terms", "Termination"])
        """
        self.doc = doc
        self.known_standards = known_standards
        self._heading_style = None
        self._body_style = None
    
    def detect_heading_style(self) -> str:
        """
        Scan document paragraphs to find the style used for standard headings.
        
        Returns:
            Style name (e.g., 'Heading 2') or 'Heading 2' as fallback
        """
        print(f"\n=== DEBUG StyleDetector.detect_heading_style ===")
        print(f"Scanning {len(self.doc.paragraphs)} paragraphs...")
        print(f"Looking for standards: {self.known_standards}")
        
        if self._heading_style:
            print(f"Using cached heading style: '{self._heading_style}'")
            return self._heading_style
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            # Check if paragraph text matches any known standard
            for standard in self.known_standards:
                if standard.lower() in text.lower():
                    # Found a match - use this paragraph's style
                    style_name = para.style.name if para.style else 'Heading 2'
                    self._heading_style = style_name
                    print(f"✓ Found '{standard}' at paragraph {i}")
                    print(f"  Text: '{text[:60]}...'")
                    print(f"  Style: '{style_name}'")
                    return style_name
        
        # No match found, use default
        self._heading_style = 'Heading 2'
        print(f"⚠ No matching standards found, using fallback: 'Heading 2'")
        return self._heading_style
    
    def detect_body_style(self) -> str:
        """
        Scan document to find the body style used after standard headings.
        
        Returns:
            Style name (e.g., 'Normal', 'Body Text') or 'Normal' as fallback
        """
        print(f"\n=== DEBUG StyleDetector.detect_body_style ===")
        
        if self._body_style:
            print(f"Using cached body style: '{self._body_style}'")
            return self._body_style
        
        # Find a standard heading, then get the style of the next paragraph
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            for standard in self.known_standards:
                if standard.lower() in text.lower():
                    # Found standard - check next paragraph
                    if i + 1 < len(self.doc.paragraphs):
                        next_para = self.doc.paragraphs[i + 1]
                        style_name = next_para.style.name if next_para.style else 'Normal'
                        self._body_style = style_name
                        print(f"✓ Found body text after '{standard}' at paragraph {i+1}")
                        print(f"  Text: '{next_para.text[:60]}...'")
                        print(f"  Style: '{style_name}'")
                        return style_name
        
        # No match found, use default
        self._body_style = 'Normal'
        print(f"⚠ No body text found after standards, using fallback: 'Normal'")
        return self._body_style


def append_suggested_standards(
    original_docx_path: Path,
    items: List[Dict[str, str]],
    known_standards: Optional[List[str]] = None
) -> Path:
    """
    Append suggested standards to a DOCX document with matched styling.
    
    Args:
        original_docx_path: Path to original DOCX file
        items: List of dicts with keys 'standard' (heading) and 'suggestion' (body text)
        known_standards: Optional list of standard names for style detection
    
    Returns:
        Path to temporary edited DOCX file
    
    Raises:
        FileNotFoundError: If original_docx_path doesn't exist
        ValueError: If items is empty or malformed
    """
    print(f"\n{'='*60}")
    print(f"=== DEBUG append_suggested_standards ===")
    print(f"{'='*60}")
    print(f"Original document path: {original_docx_path}")
    print(f"Document exists: {original_docx_path.exists()}")
    print(f"Number of items to append: {len(items) if items else 0}")
    
    if not original_docx_path.exists():
        print(f"✗ ERROR: File not found!")
        raise FileNotFoundError(f"Original document not found: {original_docx_path}")
    
    if not items:
        print(f"✗ ERROR: Empty items list!")
        raise ValueError("No items provided to append")
    
    # Validate items structure
    print(f"\nValidating items structure...")
    for i, item in enumerate(items):
        if 'standard' not in item or 'suggestion' not in item:
            print(f"✗ ERROR: Item {i} missing required keys: {list(item.keys())}")
            raise ValueError("Each item must have 'standard' and 'suggestion' keys")
        print(f"  Item {i+1}: '{item['standard'][:40]}...' ({len(item['suggestion'])} chars)")
    
    # Load document
    print(f"\nLoading document...")
    try:
        doc = Document(str(original_docx_path))
        print(f"✓ Document loaded: {len(doc.paragraphs)} paragraphs, {len(doc.sections)} sections")
    except Exception as e:
        print(f"✗ ERROR loading document: {e}")
        raise
    
    # Detect styles if we have known standards
    if known_standards is None:
        known_standards = [item['standard'] for item in items]
        print(f"Using item standards for detection: {known_standards}")
    
    print(f"\nDetecting document styles...")
    detector = StyleDetector(doc, known_standards)
    heading_style = detector.detect_heading_style()
    body_style = detector.detect_body_style()
    print(f"\n✓ Style detection complete:")
    print(f"  Heading style: '{heading_style}'")
    print(f"  Body style: '{body_style}'")
    
    # Add page break to start appendix on new page
    print(f"\nAdding page break before appendix...")
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        run = last_para.add_run()
        run.add_break(WD_BREAK.PAGE)
        print(f"✓ Page break added after paragraph {len(doc.paragraphs)-1}")
    
    # Add appendix heading
    print(f"Adding appendix heading...")
    try:
        appendix_heading = doc.add_heading('Appendix — Suggested Standards', level=1)
        print(f"✓ Appendix heading added with 'Heading 1' style")
    except KeyError:
        # Document doesn't have Heading 1 style, use paragraph with bold/large font
        print(f"⚠ 'Heading 1' style not found, using formatted paragraph instead")
        appendix_heading = doc.add_paragraph()
        run = appendix_heading.add_run('Appendix — Suggested Standards')
        run.bold = True
        run.font.size = Pt(16)
        print(f"✓ Appendix heading added as formatted paragraph")
    
    # Add each suggested standard
    print(f"\nAppending {len(items)} suggested standards:")
    for i, item in enumerate(items):
        print(f"  [{i+1}/{len(items)}] Adding '{item['standard']}'...")
        
        # Add standard heading
        standard_heading = doc.add_paragraph(item['standard'])
        try:
            standard_heading.style = heading_style
            print(f"    ✓ Heading applied style: '{heading_style}'")
        except KeyError:
            # Style doesn't exist, use Heading 2
            standard_heading.style = 'Heading 2'
            print(f"    ⚠ Style '{heading_style}' not found, using 'Heading 2'")
        
        # Add suggestion body
        suggestion_para = doc.add_paragraph(item['suggestion'])
        try:
            suggestion_para.style = body_style
            print(f"    ✓ Body applied style: '{body_style}'")
        except KeyError:
            # Style doesn't exist, use Normal
            suggestion_para.style = 'Normal'
            print(f"    ⚠ Style '{body_style}' not found, using 'Normal'")
        
        # Add spacing after each standard
        doc.add_paragraph()
    
    print(f"\n✓ All standards appended successfully")
    
    # Save to temporary file
    print(f"\nSaving edited document to temporary file...")
    try:
        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix='.docx',
            prefix='contract_edited_'
        )
        temp_path = Path(temp_file.name)
        temp_file.close()
        
        print(f"Temp file path: {temp_path}")
        doc.save(str(temp_path))
        
        file_size = temp_path.stat().st_size
        print(f"✓ Document saved successfully")
        print(f"  File: {temp_path.name}")
        print(f"  Size: {file_size:,} bytes")
        print(f"{'='*60}\n")
        
        return temp_path
    except Exception as e:
        print(f"✗ ERROR saving document: {e}")
        raise
